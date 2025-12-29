"""
RAG Agent - Streamlit Web Application

This is the main entry point for the RAG Agent web application.
Run this file to start the application:
    streamlit run app.py

Learning Note - What is Streamlit?
----------------------------------
Streamlit is a Python framework for building web apps quickly.
Key features:
- Pure Python (no HTML/CSS/JavaScript needed)
- Automatic UI updates when data changes
- Built-in widgets (file uploaders, chat interfaces, etc.)
- Perfect for data science and ML applications

How Streamlit works:
1. Your script runs from top to bottom
2. User interacts with a widget
3. Entire script reruns with new widget values
4. UI updates automatically

This is why we use st.session_state to persist data between reruns!
"""

import streamlit as st
from pathlib import Path
from datetime import datetime
import re
from io import BytesIO
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI

# Import our modules
from config import settings
from modules.document_processor import TextChunker
from modules.vector_store import ChromaDBStore
from modules.rag_query import RAGEngine
from utils.helpers import (
    validate_file_type,
    get_file_size_mb,
    format_sources,
    sanitize_filename
)

# Import research modules
from modules.web_searcher import TavilySearcher
from modules.web_content_processor import WebContentChunker
from modules.research_agent import ResearchAgent

# Initialize OpenAI client for chat functions
client = OpenAI(api_key=settings.OPENAI_API_KEY)


# ============================================================================
# Page Configuration
# ============================================================================

st.set_page_config(
    page_title="RAG Agent - Q&A & Research",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ============================================================================
# Custom CSS for better appearance
# ============================================================================

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1f77b4;
        margin-bottom: 1rem;
    }
    .source-box {
        background-color: #f0f2f6;
        border-left: 3px solid #1f77b4;
        padding: 10px;
        margin: 10px 0;
        border-radius: 5px;
    }
    .success-box {
        background-color: #d4edda;
        border-left: 3px solid #28a745;
        padding: 10px;
        margin: 10px 0;
        border-radius: 5px;
    }
    .info-box {
        background-color: #d1ecf1;
        border-left: 3px solid #17a2b8;
        padding: 10px;
        margin: 10px 0;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)


# ============================================================================
# Session State Initialization
# ============================================================================

def initialize_session_state():
    """
    Initialize all session state variables.

    Learning Note - Session State in Streamlit:
    ------------------------------------------
    Streamlit reruns your script on every interaction.
    Without session state, you'd lose all data on each rerun!

    Session state persists:
    - Between reruns (when user clicks a button)
    - For the duration of the user's session
    - Not between different users (each has their own state)

    Use it for:
    - Database connections
    - Loaded models
    - Chat history
    - Any data that should survive reruns
    """
    if 'initialized' not in st.session_state:
        # Mark as initialized
        st.session_state.initialized = True

        # Initialize vector store (expensive to create, so cache it)
        st.session_state.vector_store = ChromaDBStore()

        # Available collections (fixed list - only these three)
        st.session_state.collections = ['Bioinformatics', 'AI-Agent', 'Miscellaneous']

        # Set default collection
        st.session_state.active_collection = 'Bioinformatics'
        st.session_state.vector_store.set_collection('Bioinformatics')

        # Initialize RAG engine
        st.session_state.rag_engine = RAGEngine(st.session_state.vector_store)

        # Chat history per collection: {collection_name: [chats]}
        st.session_state.chat_history = {col: [] for col in st.session_state.collections}

        # Uploaded files per collection: {collection_name: [files]}
        st.session_state.uploaded_files = {col: [] for col in st.session_state.collections}

        # Processing status
        st.session_state.processing = False

        # App mode selection (Document Q&A or Web Research)
        st.session_state.app_mode = "📄 Document Q&A"

        # Research mode variables
        st.session_state.research_history = []
        st.session_state.current_research = None

        # Chat agent variables for Web Research
        st.session_state.research_chat_messages = []
        st.session_state.research_chat_enabled = None  # None = no mode selected yet
        st.session_state.awaiting_clarification = False
        st.session_state.awaiting_confirmation = False  # Waiting for user to confirm research
        st.session_state.pending_topic = None
        st.session_state.pending_search_params = None  # Store search parameters for confirmation


# ============================================================================
# Helper Functions
# ============================================================================

def process_uploaded_file(uploaded_file) -> dict:
    """
    Process a single uploaded file.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        dict: Processing result with success status and message

    Learning Note - File Processing Pipeline:
    1. Validate file type and size
    2. Parse document to extract text
    3. Chunk text intelligently
    4. Generate embeddings
    5. Store in vector database

    This is the document ingestion pipeline!
    """
    try:
        # Get file info
        filename = sanitize_filename(uploaded_file.name)
        file_type = Path(filename).suffix.lower()
        file_size = get_file_size_mb(uploaded_file)

        # Validate file type
        if not validate_file_type(filename):
            return {
                'success': False,
                'message': f"Unsupported file type: {file_type}. Supported: {', '.join(settings.SUPPORTED_FILE_TYPES)}"
            }

        # Validate file size
        if file_size > settings.MAX_UPLOAD_SIZE_MB:
            return {
                'success': False,
                'message': f"File too large: {file_size:.1f} MB. Maximum: {settings.MAX_UPLOAD_SIZE_MB} MB"
            }

        # Process document: parse, chunk, add metadata
        chunks = TextChunker.process_document(
            uploaded_file,
            filename,
            file_type
        )

        # Store in vector database
        count = st.session_state.vector_store.add_documents(
            chunks,
            show_progress=False
        )

        return {
            'success': True,
            'message': f"Successfully processed {filename}: {count} chunks created",
            'filename': filename,
            'chunk_count': count
        }

    except Exception as e:
        return {
            'success': False,
            'message': f"Error processing {uploaded_file.name}: {str(e)}"
        }


def display_sources(sources: list, use_expanders: bool = True):
    """
    Display retrieved sources in an attractive format.

    Args:
        sources: List of retrieved chunks with metadata
        use_expanders: Whether to use expanders for each source (set False if already in expander)

    Learning Note - Image Display:
    -----------------------------
    Sources can now include images! When a source is an image chunk:
    - Display the actual image
    - Show the vision model's description
    - Include metadata (size, page number, etc.)

    This allows users to see visual content from their documents.
    """
    if not sources:
        return

    st.markdown("### 📄 Retrieved Sources")

    for idx, source in enumerate(sources, start=1):
        metadata = source['metadata']
        filename = metadata.get('filename', 'Unknown')
        distance = source.get('distance', 0)
        chunk_type = metadata.get('chunk_type', 'text')

        # Calculate similarity score (0-1, higher is better)
        similarity = 1 / (1 + distance)

        # Check if this is an image chunk
        is_image = chunk_type == 'image'
        icon = "🖼️" if is_image else "📄"

        if use_expanders:
            expander_title = f"{icon} Source {idx}: {filename}"
            if is_image:
                # Check if it's from a slide or a page
                if 'slide_number' in metadata:
                    slide = metadata.get('slide_number', 'unknown')
                    expander_title += f" (Slide {slide})"
                else:
                    page = metadata.get('page_number', 'unknown')
                    expander_title += f" (Page {page})"
            elif 'slide_number' in metadata:
                # Text from a slide
                slide = metadata.get('slide_number', 'unknown')
                slide_title = metadata.get('slide_title', '')
                if slide_title:
                    expander_title += f" (Slide {slide}: {slide_title})"
                else:
                    expander_title += f" (Slide {slide})"
            elif 'sheet_name' in metadata:
                # Text from an Excel sheet
                sheet_name = metadata.get('sheet_name', 'unknown')
                expander_title += f" (Sheet: {sheet_name})"
            expander_title += f" - Similarity: {similarity:.2%}"

            with st.expander(expander_title):
                st.markdown(f"**File:** {filename}")
                st.markdown(f"**Type:** {'Image' if is_image else 'Text'}")
                st.markdown(f"**Relevance Score:** {similarity:.2%}")

                if is_image:
                    # Display image information
                    if 'slide_number' in metadata:
                        st.markdown(f"**Slide:** {metadata.get('slide_number', 'unknown')}")
                    else:
                        st.markdown(f"**Page:** {metadata.get('page_number', 'unknown')}")
                    st.markdown(f"**Size:** {metadata.get('image_width', '?')}x{metadata.get('image_height', '?')} pixels")

                    # Display the image
                    if 'image_data' in metadata:
                        try:
                            import base64
                            from PIL import Image as PILImage
                            from io import BytesIO

                            # Decode base64 image
                            image_bytes = base64.b64decode(metadata['image_data'])
                            image = PILImage.open(BytesIO(image_bytes))

                            st.image(image, caption=f"Image from {filename}", use_container_width=True)
                        except Exception as e:
                            st.error(f"Could not display image: {e}")

                    # Show description
                    st.markdown("**Description:**")
                    # Remove [IMAGE] prefix if present
                    description = source['text'].replace('[IMAGE] ', '')
                    st.markdown(description)
                else:
                    # Display text content
                    # Show Excel-specific metadata if available
                    if 'sheet_name' in metadata:
                        st.markdown(f"**Sheet:** {metadata.get('sheet_name', 'unknown')}")
                        if 'num_rows' in metadata and 'num_cols' in metadata:
                            st.markdown(f"**Dimensions:** {metadata['num_rows']} rows × {metadata['num_cols']} columns")
                        if 'columns' in metadata and metadata['columns']:
                            cols = metadata['columns']
                            cols_preview = cols[:5]  # Show first 5 columns
                            cols_str = ', '.join(str(c) for c in cols_preview)
                            if len(cols) > 5:
                                cols_str += f", ... ({len(cols)} total)"
                            st.markdown(f"**Columns:** {cols_str}")

                    st.markdown("**Content:**")
                    st.text(source['text'])
        else:
            # Display without expander (for when already inside an expander)
            st.markdown(f"{icon} **Source {idx}: {filename}**")
            st.markdown(f"- Type: {'Image' if is_image else 'Text'}")
            st.markdown(f"- Relevance Score: {similarity:.2%}")

            if is_image:
                st.markdown(f"- Page: {metadata.get('page_number', 'unknown')}")
                # Show thumbnail for image
                if 'image_data' in metadata:
                    try:
                        import base64
                        from PIL import Image as PILImage
                        from io import BytesIO

                        image_bytes = base64.b64decode(metadata['image_data'])
                        image = PILImage.open(BytesIO(image_bytes))

                        # Create thumbnail
                        image.thumbnail((200, 200))
                        st.image(image, width=200)
                    except:
                        pass

                description = source['text'].replace('[IMAGE] ', '')
                st.markdown(f"- Description: {description[:200]}...")
            else:
                st.markdown(f"- Content Preview: {source['text'][:200]}...")

            st.markdown("---")


def create_research_docx(research: dict) -> BytesIO:
    """
    Create a DOCX file from research report with clickable hyperlinks.

    Args:
        research: Research record dictionary with report, sources, metadata

    Returns:
        BytesIO object containing the DOCX file

    Learning Note - Why DOCX for Reports?
    ------------------------------------
    DOCX files are better than plain text because they:
    - Support clickable hyperlinks (users can verify sources)
    - Preserve formatting (headings, bold, lists)
    - Are widely compatible (Word, Google Docs, etc.)
    - Look professional
    - Can be easily edited and shared
    """
    doc = Document()

    # Add title
    title = doc.add_heading(f"Research Report: {research['topic']}", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add metadata section
    doc.add_paragraph()
    metadata_para = doc.add_paragraph()
    metadata_para.add_run("Generated: ").bold = True
    metadata_para.add_run(research['timestamp'][:19])
    metadata_para.add_run(" | ")
    metadata_para.add_run("Sources: ").bold = True
    metadata_para.add_run(str(research['num_sources']))
    metadata_para.add_run(" | ")
    metadata_para.add_run("Search Quality: ").bold = True
    metadata_para.add_run(f"{research.get('search_quality', 0.0):.2f}")

    doc.add_paragraph("_" * 80)
    doc.add_paragraph()

    # Parse and add report content with hyperlinks
    report_text = research['report']

    # Split report into lines and process
    lines = report_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            doc.add_paragraph()
            continue

        # Check for markdown headings
        if line.startswith('###'):
            doc.add_heading(line.replace('###', '').strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=1)
        else:
            # Regular paragraph - parse for hyperlinks
            para = doc.add_paragraph()
            add_hyperlinks_to_paragraph(para, line)

    # Add sources section
    doc.add_paragraph()
    doc.add_paragraph("_" * 80)
    doc.add_heading("Sources", level=1)

    for i, source in enumerate(research['sources'], 1):
        para = doc.add_paragraph(style='List Number')
        para.add_run(f"{source['title']}").bold = True
        para.add_run("\n   ")

        # Add clickable URL
        add_hyperlink(para, source['url'], source['url'])

        para.add_run(f"\n   Relevance: {source.get('relevance_score', 'N/A')}")
        para.add_run(f"\n   Retrieved: {source.get('retrieved_date', 'N/A')[:10]}")

    # Add footer
    doc.add_paragraph()
    doc.add_paragraph("_" * 80)
    footer = doc.add_paragraph("Generated by RAG Research Agent")
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def add_hyperlinks_to_paragraph(paragraph, text):
    """
    Parse text for markdown hyperlinks and add them to paragraph.

    Handles patterns like: [Source 1: Title](URL)
    """
    # Pattern to match markdown links: [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'

    last_end = 0
    for match in re.finditer(link_pattern, text):
        # Add text before the link
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        # Add the hyperlink
        link_text = match.group(1)
        url = match.group(2)
        add_hyperlink(paragraph, url, link_text)

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def add_hyperlink(paragraph, url, text):
    """
    Add a clickable hyperlink to a paragraph.

    This is a helper function because python-docx doesn't have a simple
    add_hyperlink method - we need to work with the underlying XML.
    """
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

    # Get access to the paragraph's part (for relationship management)
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run element for the hyperlink text
    new_run = OxmlElement('w:r')

    # Create run properties for styling (blue, underlined)
    rPr = OxmlElement('w:rPr')

    # Add blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    # Add the run to the hyperlink
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)


# ============================================================================
# Chat Agent Helper Functions
# ============================================================================

def quick_conversational_check(message: str) -> Dict[str, Any]:
    """
    Fast pattern-based check for common conversational queries.
    Returns response immediately without API call for common questions.

    Learning Note - Performance Optimization:
    ----------------------------------------
    API calls take 1-2 seconds. For common questions, we can skip the API
    entirely and respond instantly using pattern matching.

    This function handles:
    - Capability questions ("what can you do", "how do you work")
    - Greetings ("hi", "hello", "thanks")
    - Simple yes/no questions about features

    Only uses API for complex/ambiguous queries.
    """
    msg_lower = message.lower().strip()

    # Greeting patterns
    greeting_patterns = ['hi', 'hello', 'hey', 'greetings', 'good morning', 'good afternoon', 'good evening']
    if any(msg_lower.startswith(pattern) for pattern in greeting_patterns):
        return {
            'type': 'quick_response',
            'response': "Hello! I'm your research assistant. I can search the web and synthesize information on any topic you'd like to explore. What would you like me to research for you?"
        }

    # Thank you patterns
    thanks_patterns = ['thanks', 'thank you', 'thx', 'appreciate it', 'perfect', 'great', 'awesome', 'excellent']
    if any(pattern in msg_lower for pattern in thanks_patterns) and len(msg_lower.split()) <= 5:
        return {
            'type': 'quick_response',
            'response': "You're welcome! Let me know if you'd like to research anything else."
        }

    # Capability questions
    if any(pattern in msg_lower for pattern in ['what can you do', 'what do you do', 'your capabilities', 'what are you capable of']):
        return {
            'type': 'quick_response',
            'response': """I'm a web research assistant! Here's what I can do:

• **Research any topic** - I search multiple web sources (articles, blogs, news, social media)
• **Synthesize findings** - I analyze and combine information from different sources
• **Provide citations** - Every fact is backed by source URLs
• **Answer follow-ups** - Ask me questions about the research I've conducted

I use Tavily AI to search the web and OpenAI GPT to synthesize findings. Just tell me what topic you'd like me to research!"""
        }

    # YouTube capability questions only (not actual research requests)
    # Only match if asking about capability WITHOUT specifying a topic
    if 'youtube' in msg_lower and not any(topic_word in msg_lower for topic_word in ['about', 'on', 'topic', 'channel', 'video', 'tutorial', 'coding', 'programming']):
        if any(word in msg_lower for word in ['can you search youtube', 'do you search youtube', 'know how to search youtube']):
            return {
                'type': 'quick_response',
                'response': "Yes! I can search for information **about** YouTube topics and content. I search web sources (articles, blogs, discussions) that talk about YouTube videos, trends, and creators. I don't search YouTube's video database directly, but I can find web content that discusses YouTube topics. What would you like me to research?"
            }

    # How does it work questions
    if any(pattern in msg_lower for pattern in ['how do you work', 'how does this work', 'how does it work', 'explain how']):
        return {
            'type': 'quick_response',
            'response': """Here's how I work:

1. **You give me a topic** - Tell me what you want to research
2. **I search the web** - I use Tavily AI to search multiple sources
3. **I analyze results** - I process 15-20 sources on your topic
4. **I synthesize a report** - I combine findings into a clear, cited report
5. **You can ask follow-ups** - Ask me to clarify or dive deeper

The whole process takes 20-30 seconds. Want to try it? Give me a topic!"""
        }

    # Cost questions
    if any(pattern in msg_lower for pattern in ['how much', 'cost', 'price', 'expensive']):
        return {
            'type': 'quick_response',
            'response': "Each research session costs approximately $0.05-0.06 (about 5 cents). This includes searching multiple sources and AI synthesis. It's very affordable for comprehensive research! What would you like me to research?"
        }

    # Subscriber/metrics/analytics questions
    if any(pattern in msg_lower for pattern in ['subscriber', 'subscription', 'view count', 'views', 'analytics', 'metrics', 'statistics', 'follower']):
        if any(word in msg_lower for word in ['see', 'can you', 'do you', 'show', 'get', 'find']):
            return {
                'type': 'quick_response',
                'response': "No, I can't see subscriber counts, view counts, or live metrics directly from platforms. However, I can search for web articles that may mention this information about specific channels or topics."
            }

    # Not a quick response - needs API processing
    return {'type': 'needs_api'}


def check_topic_ambiguity(topic: str) -> Dict[str, Any]:
    """
    Check if a research topic needs clarification using GPT.

    Learning Note - Why Check for Ambiguity?
    ----------------------------------------
    Users often provide vague topics like:
    - "AI" (too broad - AI in what context?)
    - "the latest developments" (in what field?)
    - "it" (what does "it" refer to?)

    By detecting ambiguity upfront, we can ask clarifying questions
    and conduct more targeted research that actually answers what
    the user wants to know.

    Args:
        topic: The research topic entered by user

    Returns:
        Dictionary with:
        - is_ambiguous: bool (True if needs clarification)
        - clarification_questions: List[str] (questions to ask user)
        - reasoning: str (why it's ambiguous or not)

    Example:
        >>> result = check_topic_ambiguity("AI")
        >>> if result['is_ambiguous']:
        ...     for q in result['clarification_questions']:
        ...         print(f"- {q}")
    """
    try:
        # Use OpenAI to analyze the topic
        prompt = f"""You are a research assistant helping clarify user research topics.

Analyze this research topic: "{topic}"

Determine if the topic is clear enough to research, or if it's ambiguous and needs clarification.

A topic is AMBIGUOUS if:
- It's too broad or vague (e.g., "AI", "technology", "climate")
- It uses pronouns without context (e.g., "it", "this", "that")
- It's unclear what specific aspect the user wants (e.g., "latest developments" without specifying the field)
- It could mean multiple different things

A topic is CLEAR if:
- It specifies the domain and focus (e.g., "AI applications in radiology 2024")
- It's specific enough to guide research (e.g., "quantum computing error correction techniques")
- The user's intent is obvious

Respond in JSON format:
{{
    "is_ambiguous": true/false,
    "reasoning": "brief explanation of why it is/isn't ambiguous",
    "clarification_questions": ["question 1", "question 2", "question 3"] (only if ambiguous, otherwise empty list)
}}

The clarification questions should help narrow down what the user actually wants to know.
Examples:
- "Which specific application area are you interested in?"
- "Are you looking for recent developments or fundamental concepts?"
- "Which time period or year are you focusing on?"

Respond ONLY with valid JSON, no other text."""

        # Call OpenAI API
        llm_client = OpenAI(api_key=settings.OPENAI_API_KEY)
        response = llm_client.chat.completions.create(
            model="gpt-3.5-turbo",  # Use faster model for ambiguity detection
            messages=[
                {"role": "system", "content": "You are a helpful research assistant that detects ambiguous topics."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.0,  # Deterministic for consistent analysis
            max_tokens=400
        )

        # Parse response
        response_text = response.choices[0].message.content.strip()

        # Try to extract JSON from response
        import json
        import re

        # Sometimes the model wraps JSON in markdown code blocks
        json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
        if json_match:
            response_text = json_match.group(1)

        result = json.loads(response_text)

        # Validate result has required keys
        if 'is_ambiguous' not in result:
            result['is_ambiguous'] = False
        if 'reasoning' not in result:
            result['reasoning'] = 'Unable to determine'
        if 'clarification_questions' not in result:
            result['clarification_questions'] = []

        return result

    except Exception as e:
        print(f"Error checking topic ambiguity: {e}")
        # Default to not ambiguous on error (fail open)
        return {
            'is_ambiguous': False,
            'reasoning': f'Error during analysis: {str(e)}',
            'clarification_questions': []
        }


def detect_research_intent(message: str) -> Dict[str, Any]:
    """
    Detect if a message is requesting research or just conversational.

    Learning Note - Intent Detection:
    --------------------------------
    Not every message should trigger research!
    - "What can you do?" → Conversational
    - "Do you know about X?" → Conversational
    - "Research quantum computing" → Research request
    - "Tell me about AI" → Could be either (ambiguous)

    This function uses GPT to understand user intent.
    """
    try:
        # Quick keyword check for obvious research requests
        msg_lower = message.lower().strip()
        research_keywords = [
            'research', 'find information', 'search for', 'look up',
            'tell me about', 'what are', 'what is', 'how to',
            'tutorial', 'guide', 'learn about', 'popular', 'best',
            'latest', 'channels on', 'about', 'information on'
        ]

        # If message contains research keywords and is longer than 2 words, likely research
        if any(keyword in msg_lower for keyword in research_keywords) and len(message.split()) >= 2:
            # But exclude pure capability questions
            if not any(cap in msg_lower for cap in ['what can you', 'how do you work', 'how does this work']):
                return {
                    'intent': 'research_request',
                    'confidence': 0.9,
                    'reasoning': 'Contains research keywords'
                }

        # If message is 1-3 words and not a greeting/thanks, likely a topic name
        word_count = len(message.split())
        if word_count <= 3 and word_count >= 1:
            if not any(casual in msg_lower for casual in ['hi', 'hello', 'hey', 'thanks', 'thank', 'yes', 'no', 'ok', 'sure']):
                return {
                    'intent': 'research_request',
                    'confidence': 0.85,
                    'reasoning': 'Short topic name'
                }
        prompt = f"""Analyze this user message and determine their intent.

User message: "{message}"

Is the user:
A) Asking a conversational question about capabilities, features, or how things work?
   Examples: "What can you do?", "How does this work?", "Can you help me?"

B) Making a greeting or casual conversation?
   Examples: "Hello", "Thanks", "That's helpful"

C) Requesting actual research on a topic?
   Examples:
   - "Research quantum computing"
   - "Find information about AI in healthcare"
   - "vibe coding" (a topic to research)
   - "tutorial on how to vibe code"
   - "popular YouTube channels on [topic]"
   - "search for [topic]"
   - Any topic/subject name they want researched

IMPORTANT: If the message mentions a specific TOPIC or SUBJECT (like "vibe coding", "AI agents", "Python tutorials"),
it's almost certainly a research request, NOT conversational.

Respond in JSON format:
{{
    "intent": "conversational" or "research_request",
    "confidence": 0.0-1.0,
    "reasoning": "brief explanation"
}}

If unclear, prefer "research_request" over "conversational" - it's better to ask for confirmation than to ignore a research request."""

        llm_client = OpenAI(api_key=settings.OPENAI_API_KEY)
        response = llm_client.chat.completions.create(
            model="gpt-3.5-turbo",  # Use faster, cheaper model for intent detection
            messages=[
                {"role": "system", "content": "You are an intent detection system."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.0,  # Deterministic for consistent intent detection
            max_tokens=150
        )

        # Parse response
        import json
        import re
        response_text = response.choices[0].message.content.strip()

        # Extract JSON
        json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
        if json_match:
            response_text = json_match.group(1)

        result = json.loads(response_text)

        return {
            'intent': result.get('intent', 'conversational'),
            'confidence': result.get('confidence', 0.5),
            'reasoning': result.get('reasoning', '')
        }

    except Exception as e:
        print(f"Error detecting intent: {e}")
        # Default to conversational on error
        return {
            'intent': 'conversational',
            'confidence': 0.5,
            'reasoning': 'Error during detection'
        }


def generate_chat_response(user_message: str, research_context: Dict[str, Any] = None) -> str:
    """
    Generate a conversational response using the research context.

    Learning Note - Chat vs Direct Research:
    ---------------------------------------
    Direct research: User provides topic → immediate search
    Chat: User asks question → we check context → provide answer or research

    This allows:
    - Follow-up questions: "What about the cost?" (knows we're discussing quantum computing)
    - Clarifications: "I meant AI in healthcare" (after initial ambiguous query)
    - Refinements: "Focus on 2024 only" (narrowing previous results)

    Args:
        user_message: The user's chat message
        research_context: Optional dict with previous research results
            {
                'topic': str,
                'report': str,
                'sources': List[Dict],
                'chunks': List[Dict]  # Raw chunks for RAG retrieval
            }

    Returns:
        str: AI-generated response

    Example:
        >>> # After researching "quantum computing"
        >>> response = generate_chat_response(
        ...     "What about the costs?",
        ...     research_context=previous_research
        ... )
    """
    try:
        # Build context from previous research if available
        context_text = ""
        if research_context:
            context_text = f"""
Previous Research Topic: {research_context.get('topic', 'N/A')}

Previous Research Report:
{research_context.get('report', 'No previous research available')}

Use this context to answer follow-up questions. If the user's question is related to the previous research, use the information from the report. If it's a new topic, indicate that new research would be needed.
"""

        # Build conversation prompt
        prompt = f"""You are a helpful research assistant. You help users clarify their research topics and answer follow-up questions about research results.

{context_text}

User's message: {user_message}

Instructions:
- If this is a follow-up question about previous research, answer using the context provided
- If this is a new topic that needs research, acknowledge it and suggest conducting research
- Be conversational and helpful
- Keep responses concise (2-4 paragraphs max)
- If you mention sources, reference them clearly

Respond naturally to the user's message:"""

        # Call OpenAI API
        llm_client = OpenAI(api_key=settings.OPENAI_API_KEY)
        response = llm_client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are a helpful research assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Error generating chat response: {e}")
        return f"I apologize, but I encountered an error: {str(e)}. Please try rephrasing your message."


# ============================================================================
# Main Application
# ============================================================================

def main():
    """Main application function."""

    # Initialize session state
    initialize_session_state()

    # Header
    st.markdown('<h1 class="main-header">📚 RAG Agent - Q&A & Research</h1>', unsafe_allow_html=True)

    # Mode selector
    st.markdown("### Select Mode")
    mode_col1, mode_col2 = st.columns(2)
    with mode_col1:
        if st.button("📄 Document Q&A", use_container_width=True, type="primary" if st.session_state.app_mode == "📄 Document Q&A" else "secondary"):
            st.session_state.app_mode = "📄 Document Q&A"
            st.rerun()
    with mode_col2:
        if st.button("🌐 Web Research", use_container_width=True, type="primary" if st.session_state.app_mode == "🌐 Web Research" else "secondary"):
            st.session_state.app_mode = "🌐 Web Research"
            st.rerun()

    st.markdown("---")

    # Conditional rendering based on mode
    if st.session_state.app_mode == "📄 Document Q&A":
        render_document_qa_mode()
    else:
        render_research_mode()


def render_document_qa_mode():
    """Render the Document Q&A interface."""

    st.markdown("""
    Upload your documents (PDF, Word, or text files) and ask questions about them.
    The system will find relevant information and provide answers with source citations.
    """)

    # ========================================================================
    # Sidebar - Document Upload and Management
    # ========================================================================

    with st.sidebar:
        st.header("📁 Document Management")

        # Collection selector
        st.subheader("🗂️ Select Collection")
        selected_collection = st.selectbox(
            "Choose document category:",
            options=st.session_state.collections,
            index=st.session_state.collections.index(st.session_state.active_collection),
            help="Organize your documents by category"
        )

        # Switch collection if changed
        if selected_collection != st.session_state.active_collection:
            st.session_state.active_collection = selected_collection
            st.session_state.vector_store.set_collection(selected_collection)
            st.success(f"Switched to {selected_collection} collection!")
            st.rerun()

        st.markdown("---")

        # File uploader
        uploaded_files = st.file_uploader(
            "Upload documents",
            type=['pdf', 'docx', 'pptx', 'xlsx', 'xls', 'txt'],
            accept_multiple_files=True,
            help="Upload PDF, Word, PowerPoint, Excel, or text files"
        )

        # Process button
        if uploaded_files and st.button("🚀 Process Documents", type="primary"):
            st.session_state.processing = True

            with st.spinner("Processing documents..."):
                results = []

                # Process each file
                for uploaded_file in uploaded_files:
                    result = process_uploaded_file(uploaded_file)
                    results.append(result)

                # Display results
                success_count = sum(1 for r in results if r['success'])
                fail_count = len(results) - success_count

                if success_count > 0:
                    st.success(f"✅ Successfully processed {success_count} document(s)!")

                    # Add to uploaded files list for current collection
                    active_col = st.session_state.active_collection
                    for result in results:
                        if result['success']:
                            st.session_state.uploaded_files[active_col].append({
                                'filename': result['filename'],
                                'chunk_count': result['chunk_count'],
                                'upload_date': datetime.now().strftime("%Y-%m-%d %H:%M")
                            })

                if fail_count > 0:
                    st.error(f"❌ Failed to process {fail_count} document(s)")

                # Show individual results
                for result in results:
                    if result['success']:
                        st.info(f"✓ {result['message']}")
                    else:
                        st.error(f"✗ {result['message']}")

            st.session_state.processing = False

        # Display collection statistics
        st.markdown("---")
        st.subheader("📊 All Collections")

        # Show stats for only the allowed collections
        all_stats = st.session_state.vector_store.get_all_collections_stats()
        for col_stats in all_stats:
            col_name = col_stats['name']

            # Only show stats for the three allowed collections
            if col_name not in st.session_state.collections:
                continue
            doc_count = col_stats['document_count']

            # Get actual documents from ChromaDB for this collection
            # Temporarily switch to the collection to get its documents
            current_active = st.session_state.vector_store.collection_name
            st.session_state.vector_store.set_collection(col_name)
            docs_in_collection = st.session_state.vector_store.get_all_documents()
            # Switch back to the active collection
            st.session_state.vector_store.set_collection(current_active)

            # Create expander with collection name and document count
            if col_name == st.session_state.active_collection:
                expander_label = f"✓ {col_name} ({doc_count} docs) - Active"
            else:
                expander_label = f"{col_name} ({doc_count} docs)"

            with st.expander(expander_label, expanded=False):
                if docs_in_collection:
                    st.markdown("**Documents:**")
                    for i, doc in enumerate(docs_in_collection, 1):
                        st.markdown(f"{i}. 📄 {doc['filename']}")
                        st.caption(f"   Uploaded: {doc['upload_date']}")
                else:
                    st.info("No documents in this collection yet.")

        # List uploaded files for current collection
        st.markdown("---")
        active_col = st.session_state.active_collection
        uploaded_in_collection = st.session_state.uploaded_files.get(active_col, [])

        if uploaded_in_collection:
            st.subheader(f"📄 Documents in {active_col}")
            for file_info in uploaded_in_collection:
                with st.expander(f"📄 {file_info['filename']}"):
                    st.write(f"**Uploaded:** {file_info['upload_date']}")

        # Clear current collection button
        stats = st.session_state.vector_store.get_collection_stats()
        if stats['count'] > 0:
            st.markdown("---")
            if st.button(f"🗑️ Clear {active_col} Collection", help=f"Delete all documents from {active_col}"):
                if st.session_state.vector_store.clear_collection():
                    st.session_state.uploaded_files[active_col] = []
                    st.session_state.chat_history[active_col] = []
                    st.success(f"{active_col} collection cleared!")
                    st.rerun()

    # ========================================================================
    # Main Area - Chat Interface
    # ========================================================================

    # Check if documents are loaded
    stats = st.session_state.vector_store.get_collection_stats()

    if stats['count'] == 0:
        # No documents uploaded yet
        st.info("""
        👈 **Get Started:**
        1. Upload one or more documents using the sidebar
        2. Click "Process Documents"
        3. Ask questions about your documents!

        **Supported formats:** PDF, Microsoft Word (.docx), Plain Text (.txt)
        """)
    else:
        # Documents are loaded, show chat interface
        st.markdown("### 💬 Ask Questions")

        # Display chat history for current collection
        active_col = st.session_state.active_collection
        collection_chat_history = st.session_state.chat_history.get(active_col, [])

        if collection_chat_history:
            for idx, chat in enumerate(collection_chat_history):
                # Question
                with st.chat_message("user"):
                    st.write(chat['question'])

                # Answer
                with st.chat_message("assistant"):
                    st.markdown(chat['answer'])

                    # Show sources
                    if chat.get('sources'):
                        with st.expander("📚 View Sources"):
                            display_sources(chat['sources'], use_expanders=False)

        # Chat input
        question = st.chat_input("Ask a question about your documents...")

        if question:
            # Display user question
            with st.chat_message("user"):
                st.write(question)

            # Generate answer
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    try:
                        # Query the RAG engine
                        result = st.session_state.rag_engine.query(
                            question,
                            n_chunks=settings.DEFAULT_RETRIEVAL_COUNT,
                            return_sources=True
                        )

                        # Display answer
                        st.markdown(result['answer'])

                        # Display sources
                        if result.get('sources'):
                            with st.expander("📚 View Sources"):
                                display_sources(result['sources'], use_expanders=False)

                        # Add to chat history for current collection
                        active_col = st.session_state.active_collection
                        st.session_state.chat_history[active_col].append({
                            'question': question,
                            'answer': result['answer'],
                            'sources': result.get('sources', [])
                        })

                    except Exception as e:
                        st.error(f"Error generating answer: {str(e)}")
                        st.error("Please make sure your OpenAI API key is configured correctly in the .env file.")

    # ========================================================================
    # Footer
    # ========================================================================

    st.markdown("---")
    with st.expander("ℹ️ About this application"):
        st.markdown("""
        ### How it works

        This is a **RAG (Retrieval Augmented Generation)** system that:

        1. **Parses your documents** into smaller chunks
        2. **Converts text to embeddings** (numerical vectors that capture meaning)
        3. **Stores embeddings** in a vector database (ChromaDB)
        4. **Retrieves relevant chunks** when you ask a question
        5. **Generates answers** using Claude AI based on retrieved context

        ### Key Features

        - 📄 Support for PDF, Word, and text files
        - 🔍 Semantic search (finds relevant content, not just keywords)
        - 💾 Persistent storage (documents saved between sessions)
        - 📚 Source citations (see where answers come from)
        - 🤖 Powered by OpenAI's GPT models

        ### Technologies Used

        - **Streamlit**: Web interface
        - **ChromaDB**: Vector database
        - **sentence-transformers**: Text embeddings
        - **OpenAI GPT**: Answer generation
        - **pypdf & python-docx**: Document parsing

        ### Learning Resources

        - [Streamlit Documentation](https://docs.streamlit.io)
        - [ChromaDB Documentation](https://docs.trychroma.com)
        - [OpenAI Documentation](https://platform.openai.com/docs)
        - [RAG Explained](https://platform.openai.com/docs/guides/retrieval-augmented-generation)
        """)

    st.markdown("---")
    st.markdown(
        '<p style="text-align: center; color: #666;">Built with Streamlit + ChromaDB + OpenAI GPT | '
        'For educational purposes</p>',
        unsafe_allow_html=True
    )


def render_research_chat_interface():
    """
    Render the chat interface for conversational research.

    Learning Note - Why Chat for Research?
    -------------------------------------
    Chat interfaces allow:
    1. Clarification: "Did you mean AI in healthcare or AI in general?"
    2. Refinement: "Let's focus on 2024 developments only"
    3. Follow-ups: "What about the costs?" after initial research
    4. Natural interaction: Users can express intent conversationally

    This is especially helpful for:
    - Vague topics that need narrowing
    - Complex topics that benefit from dialogue
    - Users who prefer conversation over forms
    """
    st.markdown("Chat with the AI to research any topic. I'll ask clarifying questions if needed!")

    # Display chat history
    for message in st.session_state.research_chat_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Chat input
    if prompt := st.chat_input("What would you like to research?"):
        # Add user message to chat
        st.session_state.research_chat_messages.append({"role": "user", "content": prompt})

        with st.chat_message("user"):
            st.markdown(prompt)

        # Process the message
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                # Check if we're awaiting confirmation to start research
                if st.session_state.awaiting_confirmation:
                    # User responded to confirmation request
                    # Check if they confirmed (yes, proceed, go ahead, etc.) or declined (no, cancel, etc.)
                    confirmation_check = prompt.lower().strip()

                    # Simple confirmation detection
                    if any(word in confirmation_check for word in ['yes', 'yeah', 'sure', 'ok', 'okay', 'go ahead', 'proceed', 'start', 'please', 'yep', 'confirm']):
                        # User confirmed - proceed with research
                        st.session_state.awaiting_confirmation = False
                        topic = st.session_state.pending_topic
                        search_params = st.session_state.pending_search_params or {}

                        confirm_msg = f"Great! Starting research on: **{topic}**"
                        st.markdown(confirm_msg)
                        st.session_state.research_chat_messages.append({
                            "role": "assistant",
                            "content": confirm_msg
                        })

                        # Conduct research
                        try:
                            progress_placeholder = st.empty()
                            progress_placeholder.text("Starting research...")

                            research_vector_store = ChromaDBStore()
                            research_vector_store.set_collection(settings.RESEARCH_COLLECTION_NAME)
                            research_agent = ResearchAgent(
                                research_vector_store,
                                search_depth=search_params.get('depth', 'basic')
                            )

                            result = research_agent.conduct_research(
                                topic=topic,
                                num_queries=search_params.get('num_queries', 4),
                                progress_callback=lambda msg, pct: progress_placeholder.text(msg)
                            )

                            progress_placeholder.empty()

                            research_record = {
                                'topic': topic,
                                'report': result['report'],
                                'sources': result['sources'],
                                'timestamp': datetime.now().isoformat(),
                                'num_sources': len(result['sources']),
                                'num_chunks': result['num_chunks'],
                                'search_quality': result['search_quality']
                            }

                            st.session_state.current_research = research_record
                            st.session_state.research_history.insert(0, research_record)
                            st.session_state.pending_topic = None
                            st.session_state.pending_search_params = None

                            success_msg = f"✅ Research complete! I found {len(result['sources'])} sources. Check below for the full report!"
                            st.markdown(success_msg)
                            st.session_state.research_chat_messages.append({
                                "role": "assistant",
                                "content": success_msg
                            })

                            st.rerun()

                        except Exception as e:
                            error_msg = f"❌ Research failed: {str(e)}"
                            st.markdown(error_msg)
                            st.session_state.research_chat_messages.append({
                                "role": "assistant",
                                "content": error_msg
                            })
                            st.session_state.awaiting_confirmation = False
                            st.session_state.pending_topic = None
                            st.session_state.pending_search_params = None

                    elif any(word in confirmation_check for word in ['no', 'nope', 'cancel', 'stop', 'don\'t', 'nevermind', 'never mind']):
                        # User declined
                        st.session_state.awaiting_confirmation = False
                        st.session_state.pending_topic = None
                        st.session_state.pending_search_params = None

                        decline_msg = "No problem! Let me know if you'd like to research something else or if you have any questions."
                        st.markdown(decline_msg)
                        st.session_state.research_chat_messages.append({
                            "role": "assistant",
                            "content": decline_msg
                        })
                        st.rerun()

                    else:
                        # Unclear response - ask again
                        clarify_msg = "I'm not sure if you want me to proceed with the research. Could you please confirm with 'yes' or 'no'?"
                        st.markdown(clarify_msg)
                        st.session_state.research_chat_messages.append({
                            "role": "assistant",
                            "content": clarify_msg
                        })
                        st.rerun()

                # Check if we're awaiting clarification
                elif st.session_state.awaiting_clarification:
                    # User provided clarification, use it to refine the pending topic
                    refined_topic = f"{st.session_state.pending_topic} - {prompt}"
                    st.session_state.awaiting_clarification = False

                    # Ask for confirmation before starting research
                    confirm_msg = f"""Perfect! I understand you want to research: **{refined_topic}**

I'll search multiple web sources and synthesize the findings into a comprehensive report.

**Search parameters:**
- Search depth: Basic (faster)
- Number of queries: 4
- Estimated time: 20-30 seconds

Would you like me to proceed with this research? (yes/no)"""

                    st.markdown(confirm_msg)
                    st.session_state.research_chat_messages.append({
                        "role": "assistant",
                        "content": confirm_msg
                    })

                    # Mark that we're awaiting confirmation
                    st.session_state.awaiting_confirmation = True
                    st.session_state.pending_topic = refined_topic
                    st.session_state.pending_search_params = {'depth': 'basic', 'num_queries': 4}
                    st.rerun()

                else:
                    # Step 0: Quick check for common conversational queries (instant response)
                    quick_check = quick_conversational_check(prompt)

                    if quick_check['type'] == 'quick_response':
                        # Instant response without API call!
                        quick_response = quick_check['response']
                        st.markdown(quick_response)
                        st.session_state.research_chat_messages.append({
                            "role": "assistant",
                            "content": quick_response
                        })
                        st.rerun()

                    # Step 1: Detect user intent (uses fast gpt-3.5-turbo)
                    intent_result = detect_research_intent(prompt)

                    if intent_result['intent'] == 'conversational':
                        # User is asking a conversational question or greeting
                        # Respond conversationally without triggering research
                        conversational_prompt = f"""You are a helpful research assistant in chat mode.

The user said: "{prompt}"

This is a conversational question about your capabilities, not a research request.

IMPORTANT INSTRUCTIONS:
- Answer the specific question directly and honestly
- If asking about a feature you DON'T have, say "No" clearly
- If asking about a feature you DO have, say "Yes" and briefly explain
- Be concise (1-2 sentences maximum)
- Don't give generic capability overviews unless asked "what can you do"

What you CAN do:
- Search the web for information on any topic (using Tavily AI)
- Find information from websites, blogs, news, social media
- Synthesize findings into reports with source citations
- Answer follow-up questions about research

What you CANNOT do:
- See subscriber counts, view counts, or live metrics from platforms
- Access real-time analytics or statistics directly
- Search YouTube's video database directly (only find web articles about YouTube topics)
- Access content behind paywalls or login walls
- See private or restricted information

Example responses:
- "Can you see subscriber counts?" → "No, I can't see subscriber counts or live metrics. I search web articles that may mention this information."
- "Can you search YouTube?" → "I can search for information about YouTube topics from web sources, but I don't search YouTube's video database directly."
- "What can you do?" → "I'm a web research assistant. I search the web for information on any topic and synthesize findings into comprehensive reports."
"""

                        llm_client = OpenAI(api_key=settings.OPENAI_API_KEY)
                        response = llm_client.chat.completions.create(
                            model="gpt-3.5-turbo",  # Use faster model for simple conversational responses
                            messages=[
                                {"role": "system", "content": "You are a friendly research assistant."},
                                {"role": "user", "content": conversational_prompt}
                            ],
                            temperature=0.7,
                            max_tokens=200
                        )

                        conversational_response = response.choices[0].message.content.strip()
                        st.markdown(conversational_response)
                        st.session_state.research_chat_messages.append({
                            "role": "assistant",
                            "content": conversational_response
                        })
                        st.rerun()

                    else:
                        # Intent is 'research_request' - proceed with research flow

                        # Check if this is a follow-up question about existing research
                        if st.session_state.current_research and not any(
                            keyword in prompt.lower()
                            for keyword in ['research', 'study', 'investigate', 'find out about', 'tell me about', 'search for']
                        ):
                            # This looks like a follow-up question about current research
                            response = generate_chat_response(
                                prompt,
                                research_context=st.session_state.current_research
                            )
                            st.markdown(response)
                            st.session_state.research_chat_messages.append({
                                "role": "assistant",
                                "content": response
                            })
                            st.rerun()

                        else:
                            # This is a new research topic - check for ambiguity
                            ambiguity_result = check_topic_ambiguity(prompt)

                            if ambiguity_result['is_ambiguous']:
                                # Topic is ambiguous - ask for clarification
                                clarification_msg = f"""I'd be happy to research that, but I need some clarification first:

**Reasoning:** {ambiguity_result['reasoning']}

**Questions to help narrow it down:**
"""
                                for i, question in enumerate(ambiguity_result['clarification_questions'], 1):
                                    clarification_msg += f"\n{i}. {question}"

                                clarification_msg += "\n\nPlease provide more details so I can conduct better research for you!"

                                st.markdown(clarification_msg)
                                st.session_state.research_chat_messages.append({
                                    "role": "assistant",
                                    "content": clarification_msg
                                })

                                # Mark that we're awaiting clarification
                                st.session_state.awaiting_clarification = True
                                st.session_state.pending_topic = prompt
                                st.rerun()

                            else:
                                # Topic is clear - ask for confirmation before research
                                confirm_msg = f"""Great! I can research: **{prompt}**

I'll search multiple web sources and create a comprehensive report with citations.

**Search parameters:**
- Search depth: Basic (faster, good for most topics)
- Number of queries: 4 (diverse perspectives)
- Estimated time: 20-30 seconds
- Cost: ~$0.05

Would you like me to proceed with this research? (yes/no)"""

                                st.markdown(confirm_msg)
                                st.session_state.research_chat_messages.append({
                                    "role": "assistant",
                                    "content": confirm_msg
                                })

                                # Mark that we're awaiting confirmation
                                st.session_state.awaiting_confirmation = True
                                st.session_state.pending_topic = prompt
                                st.session_state.pending_search_params = {'depth': 'basic', 'num_queries': 4}
                                st.rerun()


def render_research_mode():
    """
    Render the Web Research interface.

    Learning Note - Research Interface Design:
    -----------------------------------------
    Research mode is different from Q&A:
    - Takes longer (20-30 seconds vs 2-3 seconds)
    - Needs progress indicators
    - Generates comprehensive reports vs quick answers
    - Requires different UI patterns
    """
    st.markdown("""
    Enter any topic and the AI will research it for you! The system will search multiple web sources,
    analyze findings, and generate a comprehensive research report with citations.
    """)

    # Research Mode Toggle
    st.markdown("### Select Research Mode")
    col_mode1, col_mode2 = st.columns(2)
    with col_mode1:
        if st.button("💬 Chat Mode", use_container_width=True,
                    type="primary" if st.session_state.research_chat_enabled == True else "secondary",
                    help="Conversational interface with clarification and follow-up questions"):
            st.session_state.research_chat_enabled = True
            st.rerun()
    with col_mode2:
        if st.button("🚀 Direct Research", use_container_width=True,
                    type="primary" if st.session_state.research_chat_enabled == False else "secondary",
                    help="Immediate research without conversation"):
            st.session_state.research_chat_enabled = False
            st.rerun()

    # Check if Tavily API key is configured
    if not settings.TAVILY_API_KEY:
        st.error("❌ **Tavily API key not configured**")
        st.info("""
        To use Web Research mode:
        1. Get your free API key from [tavily.com](https://tavily.com)
        2. Add it to your `.env` file: `TAVILY_API_KEY=your_key_here`
        3. Restart the application
        """)
        return

    # Sidebar - Research History
    with st.sidebar:
        st.header("📊 Research History")

        if st.session_state.research_history:
            st.markdown(f"**{len(st.session_state.research_history)} past research(es)**")

            for i, research in enumerate(st.session_state.research_history[:5]):
                with st.expander(f"📌 {research['topic'][:30]}..."):
                    if 'collection_title' in research:
                        st.caption(f"📁 {research['collection_title']}")
                    st.caption(f"Date: {research['timestamp'][:10]}")
                    st.caption(f"Sources: {research['num_sources']}")
                    if st.button(f"View", key=f"view_research_{i}"):
                        st.session_state.current_research = research
                        st.rerun()

            if len(st.session_state.research_history) > 5:
                st.caption(f"...and {len(st.session_state.research_history) - 5} more")

            st.markdown("---")
            if st.button("🗑️ Clear History"):
                st.session_state.research_history = []
                st.session_state.current_research = None
                st.success("History cleared!")
                st.rerun()
        else:
            st.info("No research history yet. Start your first research!")

    # Main area - Render appropriate interface based on mode
    if st.session_state.research_chat_enabled is None:
        # No mode selected yet - show prompt
        st.info("👆 Please select a research mode above to get started!")
    elif st.session_state.research_chat_enabled == True:
        # ===== CHAT MODE =====
        render_research_chat_interface()
    elif st.session_state.research_chat_enabled == False:
        # ===== DIRECT RESEARCH MODE =====
        # Research topic input
        topic = st.text_area(
            "Research Topic:",
            placeholder="e.g., 'Latest developments in quantum computing' or 'AI applications in healthcare 2024'",
            height=100,
            help="Enter any topic you'd like to research. The agent will search multiple sources and synthesize findings."
        )

        # Research configuration
        col1, col2 = st.columns(2)
        with col1:
            search_depth = st.selectbox(
                "Search Depth:",
                options=["basic", "advanced"],
                help="Basic: Faster and cheaper (1-2s per query). Advanced: More thorough (3-4s per query)."
            )
        with col2:
            num_queries = st.slider(
                "Number of Search Queries:",
                min_value=3,
                max_value=5,
                value=4,
                help="More queries = better coverage but higher cost. Recommended: 4"
            )

        # Start research button
        if st.button("🔍 Start Research", type="primary", disabled=not topic.strip()):
            with st.spinner("Conducting research..."):
                try:
                    # Progress tracking
                    progress_bar = st.progress(0)
                    status_text = st.empty()

                    def update_progress(msg, pct):
                        """Progress callback for research agent."""
                        status_text.text(msg)
                        progress_bar.progress(pct / 100)

                    # Initialize research components
                    status_text.text("Initializing research agent...")
                    progress_bar.progress(0.1)

                    # Use fixed collection name for all web research
                    research_vector_store = ChromaDBStore()
                    research_vector_store.set_collection(settings.RESEARCH_COLLECTION_NAME)
                    research_agent = ResearchAgent(research_vector_store, search_depth=search_depth)

                    # Conduct research
                    result = research_agent.conduct_research(
                        topic=topic,
                        num_queries=num_queries,
                        progress_callback=update_progress
                    )

                    # Clear progress indicators
                    progress_bar.empty()
                    status_text.empty()

                    # Store in session state
                    research_record = {
                        'topic': topic,
                        'report': result['report'],
                        'sources': result['sources'],
                        'timestamp': datetime.now().isoformat(),
                        'num_sources': len(result['sources']),
                        'num_chunks': result['num_chunks'],
                        'search_quality': result['search_quality']
                    }

                    st.session_state.current_research = research_record
                    st.session_state.research_history.insert(0, research_record)

                    st.success(f"✅ Research complete! Synthesized findings from {len(result['sources'])} sources.")

                    # Trigger rerun to display the report
                    st.rerun()

                except Exception as e:
                    st.error(f"❌ Research failed: {str(e)}")
                    st.error("Please check your API keys and try again.")

    # Display current or most recent research
    if st.session_state.current_research:
        research = st.session_state.current_research

        st.markdown("---")
        st.markdown(f"## Research Report: {research['topic']}")

        # Collection info
        if 'collection_title' in research:
            st.info(f"📁 Collection: **{research['collection_title']}** (`{research.get('collection_name', 'N/A')}`)")

        # Metadata
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Sources", research['num_sources'])
        with col2:
            st.metric("Chunks Processed", research.get('num_chunks', 'N/A'))
        with col3:
            quality = research.get('search_quality', 0.0)
            st.metric("Search Quality", f"{quality:.2f}")

        st.caption(f"Generated: {research['timestamp'][:10]}")

        # Display report content
        st.markdown(research['report'])

        # Display sources in expandable section
        st.markdown("---")
        with st.expander(f"📚 View All Sources ({len(research['sources'])})", expanded=False):
            for i, source in enumerate(research['sources'], 1):
                st.markdown(f"""
                **{i}. [{source['title']}]({source['url']})**
                - Relevance: {source.get('relevance_score', 'N/A')}
                - Retrieved: {source.get('retrieved_date', '')[:10]}
                """)
                if i < len(research['sources']):
                    st.markdown("---")

        # Export options
        st.markdown("---")

        # Prepare report content for download
        download_content = f"""{'=' * 80}
RESEARCH REPORT
{'=' * 80}

Topic: {research['topic']}
Generated: {research['timestamp']}
Sources: {research['num_sources']}
Chunks Processed: {research.get('num_chunks', 'N/A')}
Search Quality: {research.get('search_quality', 0.0):.2f}

{'=' * 80}

{research['report']}

{'=' * 80}
SOURCES
{'=' * 80}

"""
        # Add all sources to download content
        for i, source in enumerate(research['sources'], 1):
            download_content += f"""
{i}. {source['title']}
   URL: {source['url']}
   Relevance Score: {source.get('relevance_score', 'N/A')}
   Retrieved: {source.get('retrieved_date', 'N/A')[:10]}

"""

        download_content += f"\n{'=' * 80}\n"
        download_content += "Generated by RAG Research Agent\n"
        download_content += f"{'=' * 80}\n"

        # Create filename based on topic and timestamp
        filename_base = f"research_{research['topic'][:30].replace(' ', '_')}_{research['timestamp'][:10]}"
        txt_filename = f"{filename_base}.txt"
        docx_filename = f"{filename_base}.docx"

        # Create two columns for download buttons
        col1, col2 = st.columns(2)

        with col1:
            st.download_button(
                label="📥 Download as Text File",
                data=download_content,
                file_name=txt_filename,
                mime="text/plain",
                help="Plain text format - no hyperlinks"
            )

        with col2:
            # Generate DOCX file
            docx_file = create_research_docx(research)

            st.download_button(
                label="📥 Download as DOCX (Recommended)",
                data=docx_file,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                help="Word document with clickable hyperlinks"
            )


# ============================================================================
# Entry Point
# ============================================================================

if __name__ == "__main__":
    # This block runs when you execute: streamlit run app.py
    main()
